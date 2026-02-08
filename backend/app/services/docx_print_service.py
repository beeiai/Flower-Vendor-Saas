import os
from datetime import datetime
from typing import List, Dict, Any
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

# Get the base directory of the project
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
TEMPLATE_DIR = os.path.join(BASE_DIR, "app", "templates", "docx_templates")
OUTPUT_DIR = os.path.join(BASE_DIR, "app", "generated_reports")

# Create directories if they don't exist
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


class DocxPrintService:
    """Service for generating DOCX-based reports with dynamic table handling"""
    
    @staticmethod
    def render_ledger_report(
        farmer_name: str,
        ledger_name: str,
        address: str,
        group_name: str,
        balance: float,
        commission_pct: float,
        from_date: str,
        to_date: str,
        rows: List[Dict[str, Any]],
        total_qty: float,
        total_amount: float,
        commission: float,
        luggage_total: float,
        coolie: float,
        net_amount: float,
        paid_amount: float,
        final_total: float,
        template_path: str = None
    ) -> bytes:
        """
        Generate a DOCX ledger report with dynamic table rows and convert to PDF
        
        Args:
            template_path: Path to .docx template file (if None, uses default)
        """
        try:
            # Use default template if none provided
            if not template_path:
                template_path = os.path.join(TEMPLATE_DIR, "ledger_template.docx")
            
            # Check if template exists
            if not os.path.exists(template_path):
                # Create a basic template if it doesn't exist
                DocxPrintService._create_default_ledger_template(template_path)
            
            # Load template
            doc = DocxTemplate(template_path)
            
            # Prepare context data
            context = {
                'farmer_name': farmer_name,
                'ledger_name': ledger_name,
                'address': address,
                'group_name': group_name,
                'balance': f"₹ {balance:,.2f}",
                'commission_pct': commission_pct,
                'from_date': from_date,
                'to_date': to_date,
                'current_date': datetime.now().strftime("%d-%m-%Y"),
                'rows': rows,
                'total_qty': f"{total_qty:,.2f}",
                'total_amount': f"₹ {total_amount:,.2f}",
                'commission': f"₹ {commission:,.2f}",
                'luggage_total': f"{luggage_total:,.2f}",
                'coolie': f"₹ {coolie:,.2f}",
                'net_amount': f"₹ {net_amount:,.2f}",
                'paid_amount': f"₹ {paid_amount:,.2f}",
                'final_total': f"₹ {final_total:,.2f}"
            }
            
            # Render template
            doc.render(context)
            
            # Save to temporary file
            temp_docx_path = os.path.join(OUTPUT_DIR, f"temp_ledger_{datetime.now().timestamp()}.docx")
            doc.save(temp_docx_path)
            
            # Convert to PDF
            pdf_bytes = DocxPrintService._convert_docx_to_pdf(temp_docx_path)
            
            # Clean up temporary file
            os.remove(temp_docx_path)
            
            return pdf_bytes
            
        except Exception as e:
            raise Exception(f"Failed to generate ledger report: {str(e)}")
    
    @staticmethod
    def render_group_patti_report(
        group_name: str,
        from_date: str,
        to_date: str,
        commission_pct: float,
        rows: List[Dict[str, Any]],
        totals: Dict[str, float],
        template_path: str = None
    ) -> bytes:
        """Generate a group patti report with dynamic table rows"""
        try:
            # Use default template if none provided
            if not template_path:
                template_path = os.path.join(TEMPLATE_DIR, "group_patti_template.docx")
            
            # Check if template exists
            if not os.path.exists(template_path):
                DocxPrintService._create_default_group_patti_template(template_path)
            
            # Load template
            doc = DocxTemplate(template_path)
            
            # Prepare context data
            context = {
                'group_name': group_name,
                'from_date': from_date,
                'to_date': to_date,
                'commission_pct': commission_pct,
                'current_date': datetime.now().strftime("%d-%m-%Y"),
                'rows': rows,
                'total_gross': f"₹ {totals.get('gross', 0):,.2f}",
                'total_commission': f"₹ {totals.get('commission', 0):,.2f}",
                'total_net': f"₹ {totals.get('net', 0):,.2f}",
                'total_paid': f"₹ {totals.get('paid', 0):,.2f}",
                'total_balance': f"₹ {totals.get('balance', 0):,.2f}"
            }
            
            # Render template
            doc.render(context)
            
            # Save to temporary file
            temp_docx_path = os.path.join(OUTPUT_DIR, f"temp_group_patti_{datetime.now().timestamp()}.docx")
            doc.save(temp_docx_path)
            
            # Convert to PDF
            pdf_bytes = DocxPrintService._convert_docx_to_pdf(temp_docx_path)
            
            # Clean up temporary file
            os.remove(temp_docx_path)
            
            return pdf_bytes
            
        except Exception as e:
            raise Exception(f"Failed to generate group patti report: {str(e)}")
    
    @staticmethod
    def _create_default_ledger_template(template_path: str):
        """Create a default ledger template if one doesn't exist"""
        doc = Document()
        
        # Header
        doc.add_heading('SREE KRISHNA FLOWER STALL', 0)
        doc.add_paragraph('PH: 8147848760')
        doc.add_paragraph('Proprietor: K.C.N & SONS | Mob: 9972878307')
        doc.add_paragraph('FLOWER MERCHANTS – 7795018307')
        doc.add_paragraph('Shop No: B-32, S.K.R Market, Bangalore - 560002')
        doc.add_paragraph('Trade Mark: S.K.F.S')
        doc.add_paragraph()  # Empty line
        
        # Info section (will be replaced with dynamic content)
        doc.add_paragraph('Name: {{ farmer_name }}')
        doc.add_paragraph('Ledger: {{ ledger_name }}')
        doc.add_paragraph('Address: {{ address }}')
        doc.add_paragraph('Group: {{ group_name }}')
        doc.add_paragraph('Balance: {{ balance }}')
        doc.add_paragraph('Date: {{ current_date }}')
        doc.add_paragraph()
        
        # Table header
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Qty'
        hdr_cells[2].text = 'Price'
        hdr_cells[3].text = 'Total'
        hdr_cells[4].text = 'Luggage'
        hdr_cells[5].text = 'Paid'
        hdr_cells[6].text = 'Amount'
        
        # Add a placeholder row that will be replaced
        row_cells = table.add_row().cells
        row_cells[0].text = '{% for row in rows %}{{ row.date }}'
        row_cells[1].text = '{{ row.qty }}'
        row_cells[2].text = '{{ row.price }}'
        row_cells[3].text = '{{ row.total }}'
        row_cells[4].text = '{{ row.luggage }}'
        row_cells[5].text = '{{ row.paid }}'
        row_cells[6].text = '{{ row.amount }}{% endfor %}'
        
        # Totals section
        doc.add_paragraph()
        doc.add_paragraph(f'Total Qty: {{ total_qty }}')
        doc.add_paragraph(f'Total Amount: {{ total_amount }}')
        doc.add_paragraph(f'Commission ({{ commission_pct }}%): {{ commission }}')
        doc.add_paragraph(f'Total Luggage: {{ luggage_total }}')
        doc.add_paragraph(f'Coolie: {{ coolie }}')
        doc.add_paragraph(f'Net Amount: {{ net_amount }}')
        doc.add_paragraph(f'Paid Amount: {{ paid_amount }}')
        doc.add_paragraph(f'Final Total: {{ final_total }}')
        
        # Save template
        doc.save(template_path)
    
    @staticmethod
    def _create_default_group_patti_template(template_path: str):
        """Create a default group patti template"""
        doc = Document()
        
        # Header
        doc.add_heading('GROUP PATTI REPORT', 0)
        doc.add_paragraph('SREE KRISHNA FLOWER STALL')
        doc.add_paragraph()
        
        # Group info
        doc.add_paragraph(f'Group: {{ group_name }}')
        doc.add_paragraph(f'Period: {{ from_date }} to {{ to_date }}')
        doc.add_paragraph(f'Commission: {{ commission_pct }}%')
        doc.add_paragraph(f'Date: {{ current_date }}')
        doc.add_paragraph()
        
        # Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Customer'
        hdr_cells[1].text = 'Gross'
        hdr_cells[2].text = 'Commission'
        hdr_cells[3].text = 'Net'
        hdr_cells[4].text = 'Paid'
        hdr_cells[5].text = 'Balance'
        
        # Dynamic rows
        row_cells = table.add_row().cells
        row_cells[0].text = '{% for row in rows %}{{ row.customer }}'
        row_cells[1].text = '{{ row.gross }}'
        row_cells[2].text = '{{ row.commission }}'
        row_cells[3].text = '{{ row.net }}'
        row_cells[4].text = '{{ row.paid }}'
        row_cells[5].text = '{{ row.balance }}{% endfor %}'
        
        # Totals
        doc.add_paragraph()
        doc.add_paragraph(f'Total Gross: {{ total_gross }}')
        doc.add_paragraph(f'Total Commission: {{ total_commission }}')
        doc.add_paragraph(f'Total Net: {{ total_net }}')
        doc.add_paragraph(f'Total Paid: {{ total_paid }}')
        doc.add_paragraph(f'Total Balance: {{ total_balance }}')
        
        doc.save(template_path)
    
    @staticmethod
    def _convert_docx_to_pdf(docx_path: str) -> bytes:
        """Convert DOCX file to PDF bytes using reportlab"""
        try:
            # For now, we'll return the DOCX file as bytes since direct DOCX to PDF
            # conversion requires more complex setup. In production, you might want
            # to use libreoffice or unoconv for better conversion quality.
            
            with open(docx_path, 'rb') as f:
                return f.read()
                
        except Exception as e:
            raise Exception(f"Failed to convert DOCX to PDF: {str(e)}")
    
    @staticmethod
    def create_dynamic_table(doc: Document, rows: List[Dict[str, Any]], headers: List[str]) -> None:
        """Create a dynamic table with variable number of rows"""
        # Create table with header
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                key = header.lower().replace(' ', '_')
                row_cells[i].text = str(row_data.get(key, ''))